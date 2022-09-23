import PyPDF2
from nltk import sent_tokenize, word_tokenize
import torch
from transformers import BertTokenizerFast, EncoderDecoderModel, AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
import os
from datetime import datetime
import docx2txt


def read_word(path):
    text = docx2txt.process(path)
    text_sentences = sent_tokenize(text)
    return text_sentences


def read_pdf(path):
    pdf_content = []

    with open(path, "rb") as pdf_file:
        # open file with reader
        pdf = PyPDF2.PdfFileReader(pdf_file)
        number_of_pages = pdf.getNumPages()

        for i in range(0, number_of_pages):
            page = pdf.pages[i]
            page_text = page.extractText()

            # split into sentences
            sentences = sent_tokenize(page_text)
            for sentence in sentences:
                sentence_wo_linebreak = sentence.replace("\n", "")
                pdf_content.append(sentence_wo_linebreak)

    return pdf_content


def generate_summary_with_bert(text):
    inputs = tokenizer_bert([text], padding="max_length", truncation=True, max_length=512, return_tensors="pt")
    # TODO consider max_length of input; chunk up; 512 maximal mögliche Anzahl an Tokens in diesem Modell
    input_ids = inputs.input_ids.to(device)
    attention_mask = inputs.attention_mask.to(device)
    output = model_bert.generate(input_ids, attention_mask=attention_mask, max_length=1000)
    # TODO: what is a good maxlength for output?
    return tokenizer_bert.decode(output[0], skip_special_tokens=True)


def generate_summary_with_t5(text):
    inputs = tokenizer_t5.encode(text, truncation=True, max_length=512, return_tensors="pt")
    summary_ids = model_t5.generate(inputs, max_length=150, min_length=30, length_penalty=5., num_beams=2)
    # min/max length => min/max number of tokens in summary
    # length_penalty => penalize model more or less for producing summary above/below threshold
    # num_beams => number of beams that explore the potential tokens for most promising predictions ?!
    return tokenizer_t5.decode(summary_ids[0], skip_special_tokens=True)


def chunk_up_sentences(sentence_list, max_chunk_length):
    chunk_list = []
    current_index = 0

    while current_index < len(sentence_list):
        current_chunk = ""
        current_chunk += sentence_list[current_index]
        current_chunk += " "
        current_index += 1

        # add another sentence to current chunk as long as new word length will be <= max_chunk_length
        word_length = len(word_tokenize(current_chunk))
        if current_index >= len(sentence_list):
            break
        additional_word_length = len(word_tokenize(sentence_list[current_index]))

        while word_length + additional_word_length <= max_chunk_length:
            current_chunk += sentence_list[current_index]
            current_chunk += " "
            current_index += 1
            word_length += additional_word_length
            if current_index >= len(sentence_list):
                break
            additional_word_length = len(word_tokenize(sentence_list[current_index]))

        chunk_list.append(current_chunk)

    return chunk_list


def write_word_output_file(text_summary_list_bert, text_list, file, text_summary_list_t5):
    document = Document()
    document.add_heading('IT-Camp 2022 Summarization MVP', level=0)

    document.add_heading("Zusammenfassung - BERT Model:", level=1)
    document.add_paragraph('')
    for text_summary in text_summary_list_bert:
        document.add_paragraph(text_summary)

    document.add_heading("Zusammenfassung - T5 Model:", level=1)
    document.add_paragraph('')
    for text_summary in text_summary_list_t5:
        document.add_paragraph(text_summary)

    document.add_heading("Gesamter Text:", level=1)
    for text in text_list:
        document.add_paragraph('')
        document.add_paragraph(text)

    document.save(file)


def get_current_time():
    now = datetime.now()
    return now.strftime("%H:%M:%S")


print("Started: " + get_current_time())

# 0. initialize model
# BERT
device = 'cuda' if torch.cuda.is_available() else 'cpu'
pretrained_bert = 'mrm8488/bert2bert_shared-german-finetuned-summarization'
tokenizer_bert = BertTokenizerFast.from_pretrained(pretrained_bert)
model_bert = EncoderDecoderModel.from_pretrained(pretrained_bert).to(device)
print("BERT Model & Tokenizer ready: " + get_current_time())
# T5
pretrained_t5 = 'Einmalumdiewelt/T5-Base_GNAD'
tokenizer_t5 = AutoTokenizer.from_pretrained(pretrained_t5)
model_t5 = AutoModelForSeq2SeqLM.from_pretrained(pretrained_t5, return_dict=True)
print("T5 Model & Tokenizer ready: " + get_current_time())

# 1. get interview text from file
sentences = read_word("test_files/weidmann_interview_2.docx") # TODO File Verzeichnis entsprechend anpassen
print("Read sentences: " + get_current_time())

# 2. get chunks of sentences
# TODO: divide by logical paragraphs; use model to group up sentences? filter out questions?
chunks = chunk_up_sentences(sentences, 300)
print("Chunked up sentences: " + get_current_time())

# 3. generate summary from content
# with BERT Model
summary_list_bert = []
length = str(len(chunks))
print("Summarizing with BERT Model...")
for i, chunk in enumerate(chunks):
    summary = generate_summary_with_bert(chunk)
    summary_list_bert.append(summary)
    index = str(i+1)
    print("     Generated summary " + index + "/" + length + ": " + get_current_time())
# with T5 Model
summary_list_t5 = []
print("Summarizing with T5 Model...")
for i, chunk in enumerate(chunks):
    summary = generate_summary_with_t5(chunk)
    summary_list_t5.append(summary)
    index = str(i+1)
    print("     Generated summary " + index + "/" + length + ": " + get_current_time())

# 4. generate output
filename = "auto_generated_summary.docx"
write_word_output_file(summary_list_bert, chunks, filename, summary_list_t5)
print("Created word file: " + get_current_time())

# 5. open auto-generated word file
os.startfile(filename)
